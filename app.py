import streamlit as st
import anthropic
import json
import os
import datetime
from pypdf import PdfReader
from io import BytesIO

from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import mm

from docx import Document
from docx.shared import Pt, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Page config ──────────────────────────────────────────────────────────────
st.set_page_config(page_title="AI Recruiting Workflow", layout="wide")
st.title("AI Candidate Matching & Recommendation")

# ── Sidebar: API Key ──────────────────────────────────────────────────────────
with st.sidebar:
    st.header("Settings")
    api_key = st.text_input("Anthropic API Key", type="password")

# ── Helper functions ──────────────────────────────────────────────────────────
def get_client(key):
    return anthropic.Anthropic(api_key=key.strip())

def ask(client, prompt, system=None):
    kwargs = dict(
        model="claude-sonnet-4-6",
        max_tokens=1000,
        temperature=0,
        messages=[{"role": "user", "content": prompt}]
    )
    if system:
        kwargs["system"] = system
    msg = client.messages.create(**kwargs)
    return msg.content[0].text

def ask_json(client, prompt, system=None):
    raw = ask(client, prompt, system=system)
    clean = raw.strip()
    if clean.startswith("```"):
        clean = clean.split("\n", 1)[-1]
        clean = clean.rsplit("```", 1)[0]
    return json.loads(clean.strip())

def extract_pdf_text(file):
    reader = PdfReader(file)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text.encode("utf-8", errors="ignore").decode("utf-8")

def safe(text):
    return str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

# ── File upload ───────────────────────────────────────────────────────────────
col1, col2 = st.columns(2)

with col1:
    st.subheader("Job Description")
    jd_file = st.file_uploader("Upload JD PDF", type="pdf")

with col2:
    st.subheader("Candidate Resumes")
    resume_files = st.file_uploader("Upload Resume PDFs", type="pdf", accept_multiple_files=True)

# ── Run button ────────────────────────────────────────────────────────────────
run = st.button("Run Evaluation", type="primary", disabled=not (api_key and jd_file and resume_files))

if api_key:
    st.session_state["api_key"] = api_key

if run:
    client = get_client(api_key)

    # Extract JD text
    jd_text = extract_pdf_text(jd_file)

    # Extract resume texts
    resumes = {}
    for f in resume_files:
        name = os.path.splitext(f.name)[0]
        resumes[name] = extract_pdf_text(f)

    # ── Step 1: Parse JD ─────────────────────────────────────────────────────
    with st.spinner("Parsing JD requirements..."):
        jd_parse_prompt = f'''You are a senior recruiting consultant at an AI-native recruiting firm.

Analyze the following job description and extract structured requirements.
Output ONLY valid JSON, no preamble, no markdown backticks.

Job Description:
{jd_text}

Return this exact JSON structure:
{{
  "must_have": ["list of non-negotiable requirements"],
  "strong_signals": ["list of differentiating qualities that significantly boost recommendation"],
  "nice_to_have": ["list of bonus qualifications that slightly boost recommendation"],
  "dealbreakers": ["list of conditions that immediately disqualify a candidate"],
  "work_style": ["list of behavioral and soft skill requirements"]
}}'''
        jd_requirements = ask_json(client, jd_parse_prompt)

    # ── Step 2: Evaluate candidates ───────────────────────────────────────────
    SYSTEM_PROMPT = '''You are a senior recruiting consultant at an AI-native recruiting firm.
You evaluate candidates for technical AI product roles at early-stage startups.
Be specific, evidence-based, and practical in your assessments.
Do not ask clarifying questions. Provide complete output directly.
Always output only valid JSON, no preamble, no markdown backticks.'''

    results = []
    progress = st.progress(0)
    status = st.empty()

    for i, (name, resume_text) in enumerate(resumes.items()):
        status.text(f"Evaluating {name}...")

        eval_prompt = f'''Evaluate this candidate against the following role requirements.

ROLE REQUIREMENTS:
{json.dumps(jd_requirements, indent=2)}

CANDIDATE RESUME:
{resume_text}

Return this exact JSON structure:
{{
  "candidate_name": "full name extracted from resume",
  "recommendation": "Move Forward | Validate First | Pass",
  "brief_reason": "one sentence explaining the recommendation",
  "main_risk": "one sentence describing the main risk",
  "strengths": ["strength with specific evidence from resume"],
  "risks": ["risk with specific evidence or absence of evidence"]
}}'''

        evaluation = ask_json(client, eval_prompt, system=SYSTEM_PROMPT)
        recommendation = evaluation["recommendation"]

        validation_questions = []
        candidate_summary = ""

        if recommendation != "Pass":
            depth = "1-2 lightweight confirmation questions" if recommendation == "Move Forward" else "3-4 in-depth questions targeting specific risks"
            validation_prompt = f'''Based on this candidate evaluation, generate recruiter validation questions.

CANDIDATE: {name}
EVALUATION:
{json.dumps(evaluation, indent=2)}

Generate {depth} a recruiter should ask this candidate before submission.
Questions should target unverified claims, gaps, or risks.
Be specific - reference actual items from the resume.

Return this exact JSON structure:
{{
  "questions": ["question 1", "question 2"]
}}'''
            validation_data = ask_json(client, validation_prompt, system=SYSTEM_PROMPT)
            validation_questions = validation_data["questions"]

            summary_prompt = f'''Write a client-ready candidate summary for submission to a hiring manager.

CANDIDATE: {name}
EVALUATION:
{json.dumps(evaluation, indent=2)}

Write 2-3 sentences in professional business language suitable for a client hiring manager.
Focus on what makes this candidate compelling for the role.
Do not use technical jargon. Do not mention scores or internal evaluation details.

Return this exact JSON structure:
{{
  "summary": "2-3 sentence summary here"
}}'''
            summary_data = ask_json(client, summary_prompt, system=SYSTEM_PROMPT)
            candidate_summary = summary_data["summary"]

        results.append({
            "name": f"{evaluation.get('candidate_name', name)} ({name})",
            "recommendation": recommendation,
            "brief_reason": evaluation["brief_reason"],
            "main_risk": evaluation["main_risk"],
            "strengths": evaluation["strengths"],
            "risks": evaluation["risks"],
            "validation_questions": validation_questions,
            "candidate_summary": candidate_summary
        })

        progress.progress((i + 1) / len(resumes))

    status.text("All candidates evaluated.")

    order = {"Move Forward": 0, "Validate First": 1, "Pass": 2}
    results.sort(key=lambda x: order.get(x["recommendation"], 1))

    # Store results in session state
    st.session_state["results"] = results

# ── Display results ───────────────────────────────────────────────────────────
if "results" in st.session_state:
    results = st.session_state["results"]

    st.divider()
    st.subheader("Recommendation Summary")

    REC_COLORS = {
        "Move Forward": "green",
        "Validate First": "orange",
        "Pass": "red"
    }

    # Summary table
    table_data = []
    for r in results:
        table_data.append({
            "Candidate": r["name"],
            "Recommendation": r["recommendation"],
            "Brief Reason": r["brief_reason"],
            "Main Risk": r["main_risk"]
        })
    st.dataframe(table_data, use_container_width=True)

    # Detailed reports
    st.divider()
    st.subheader("Detailed Candidate Reports")

    for idx, r in enumerate(results):
        color = REC_COLORS.get(r["recommendation"], "gray")
        with st.expander(f"{r['name']} — :{color}[{r['recommendation']}]"):
            st.markdown(f"**Reason:** {r['brief_reason']}")
            st.markdown(f"**Main Risk:** {r['main_risk']}")

            st.markdown("**Strengths**")
            for s in r["strengths"]:
                st.markdown(f"- {s}")

            st.markdown("**Risks**")
            for s in r["risks"]:
                st.markdown(f"- {s}")

            if r["validation_questions"]:
                st.markdown("**Validation Questions**")
                vq_edit_key = f"vq_edit_{idx}"
                questions_text = "\n".join([f"{i}. {q}" for i, q in enumerate(r["validation_questions"], 1)])

                if st.session_state.get(vq_edit_key, False):
                    edited_questions = st.text_area(
                        "Edit validation questions",
                        value=questions_text,
                        key=f"vq_input_{idx}",
                        label_visibility="collapsed"
                    )
                    col_save, col_cancel, _ = st.columns([1, 1, 6])
                    if col_save.button("Save", key=f"vq_save_{idx}"):
                        st.session_state["results"][idx]["validation_questions"] = [
                            q.split(". ", 1)[-1] for q in edited_questions.strip().split("\n") if q.strip()
                        ]
                        st.session_state[vq_edit_key] = False
                        st.rerun()
                    if col_cancel.button("Cancel", key=f"vq_cancel_{idx}"):
                        st.session_state[vq_edit_key] = False
                        st.rerun()
                else:
                    for i, q in enumerate(r["validation_questions"], 1):
                        st.markdown(f"{i}. {q}")
                    if st.button("Edit", key=f"vq_edit_btn_{idx}"):
                        st.session_state[vq_edit_key] = True
                        st.rerun()

            if r["candidate_summary"]:
                st.markdown("**Candidate Summary**")
                cs_edit_key = f"cs_edit_{idx}"

                if st.session_state.get(cs_edit_key, False):
                    edited_summary = st.text_area(
                        "Edit candidate summary",
                        value=r["candidate_summary"],
                        key=f"cs_input_{idx}",
                        label_visibility="collapsed"
                    )
                    col_save, col_cancel, _ = st.columns([1, 1, 6])
                    if col_save.button("Save", key=f"cs_save_{idx}"):
                        st.session_state["results"][idx]["candidate_summary"] = edited_summary
                        st.session_state[cs_edit_key] = False
                        st.rerun()
                    if col_cancel.button("Cancel", key=f"cs_cancel_{idx}"):
                        st.session_state[cs_edit_key] = False
                        st.rerun()
                else:
                    st.info(r["candidate_summary"])
                    if st.button("Edit", key=f"cs_edit_btn_{idx}"):
                        st.session_state[cs_edit_key] = True
                        st.rerun()

            # Manual trigger for Pass
            if r["recommendation"] == "Pass":
                st.divider()
                override_key = f"override_{idx}"
                if st.button("Generate Validation Questions & Summary", key=f"btn_{idx}"):
                    saved_key = st.session_state.get("api_key", "")
                    if not saved_key:
                        st.error("Please enter your API key in the sidebar.")
                    else:
                        override_client = get_client(saved_key)
                        with st.spinner("Generating..."):
                            SYSTEM_PROMPT = '''You are a senior recruiting consultant at an AI-native recruiting firm.
You evaluate candidates for technical AI product roles at early-stage startups.
Be specific, evidence-based, and practical in your assessments.
Do not ask clarifying questions. Provide complete output directly.
Always output only valid JSON, no preamble, no markdown backticks.'''

                            validation_prompt = f'''Based on this candidate evaluation, generate recruiter validation questions.

CANDIDATE: {r["name"]}
EVALUATION:
{json.dumps({"recommendation": r["recommendation"], "brief_reason": r["brief_reason"], "main_risk": r["main_risk"], "strengths": r["strengths"], "risks": r["risks"]}, indent=2)}

Generate 3-4 in-depth questions targeting specific risks a recruiter should ask this candidate.
Questions should target unverified claims, gaps, or risks.
Be specific - reference actual items from the resume.

Return this exact JSON structure:
{{
  "questions": ["question 1", "question 2"]
}}'''
                            validation_data = ask_json(override_client, validation_prompt, system=SYSTEM_PROMPT)

                            summary_prompt = f'''Write a client-ready candidate summary for submission to a hiring manager.

CANDIDATE: {r["name"]}
EVALUATION:
{json.dumps({"recommendation": r["recommendation"], "brief_reason": r["brief_reason"], "main_risk": r["main_risk"], "strengths": r["strengths"], "risks": r["risks"]}, indent=2)}

Write 2-3 sentences in professional business language suitable for a client hiring manager.
Focus on what makes this candidate compelling for the role.
Do not use technical jargon. Do not mention scores or internal evaluation details.

Return this exact JSON structure:
{{
  "summary": "2-3 sentence summary here"
}}'''
                            summary_data = ask_json(override_client, summary_prompt, system=SYSTEM_PROMPT)

                            st.session_state[override_key] = {
                                "questions": validation_data["questions"],
                                "summary": summary_data["summary"]
                            }

                if override_key in st.session_state:
                    override = st.session_state[override_key]
                    st.markdown("**Validation Questions** *(manually generated)*")
                    for i, q in enumerate(override["questions"], 1):
                        st.markdown(f"{i}. {q}")
                    st.markdown("**Candidate Summary** *(manually generated)*")
                    st.info(override["summary"])

    # ── Export PDF ────────────────────────────────────────────────────────────
    st.divider()

    def generate_pdf(results):
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4,
            rightMargin=20*mm, leftMargin=20*mm, topMargin=20*mm, bottomMargin=20*mm)

        styles = getSampleStyleSheet()
        s_title = ParagraphStyle("T", parent=styles["Title"], fontSize=16, spaceAfter=6)
        s_sub = ParagraphStyle("S", parent=styles["Normal"], fontSize=9, textColor=colors.grey, spaceAfter=12)
        s_h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontSize=12, textColor=colors.HexColor("#1a1a2e"), spaceAfter=4)
        s_h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=10, textColor=colors.HexColor("#2c5282"), spaceAfter=3)
        s_body = ParagraphStyle("B", parent=styles["Normal"], fontSize=9, leading=14, spaceAfter=4)

        REC_COLORS_PDF = {
            "Move Forward": colors.HexColor("#276749"),
            "Validate First": colors.HexColor("#744210"),
            "Pass": colors.HexColor("#742a2a")
        }

        story = []
        story.append(Spacer(1, 8*mm))
        story.append(Paragraph("Candidate Matching & Recommendation Report", s_title))
        story.append(Paragraph(
            f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')} | Candidates: {len(results)}",
            s_sub))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#e0e0e0")))
        story.append(Spacer(1, 6*mm))

        story.append(Paragraph("Recommendation Summary", s_h1))
        story.append(Spacer(1, 3*mm))
        s_header = ParagraphStyle("H", parent=styles["Normal"], fontSize=8, textColor=colors.white, fontName="Helvetica-Bold")
        table_data = [[
            Paragraph("Candidate", s_header),
            Paragraph("Recommendation", s_header),
            Paragraph("Brief Reason", s_header),
            Paragraph("Main Risk", s_header)
        ]]
        for r in results:
            table_data.append([
                Paragraph(safe(r["name"]), s_body),
                Paragraph(safe(r["recommendation"]), s_body),
                Paragraph(safe(r["brief_reason"]), s_body),
                Paragraph(safe(r["main_risk"]), s_body)
            ])
        t = Table(table_data, colWidths=[30*mm, 28*mm, 55*mm, 55*mm])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#2c5282")),
            ("TEXTCOLOR", (0,0), (-1,0), colors.white),
            ("FONTSIZE", (0,0), (-1,-1), 8),
            ("PADDING", (0,0), (-1,-1), 5),
            ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, colors.HexColor("#f7fafc")]),
            ("GRID", (0,0), (-1,-1), 0.5, colors.HexColor("#e0e0e0")),
            ("VALIGN", (0,0), (-1,-1), "TOP"),
            ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
        ]))
        story.append(t)
        story.append(Spacer(1, 8*mm))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#e0e0e0")))
        story.append(Spacer(1, 6*mm))

        story.append(Paragraph("Detailed Candidate Reports", s_h1))
        story.append(Spacer(1, 4*mm))

        for r in results:
            rec_color = REC_COLORS_PDF.get(r["recommendation"], colors.grey)
            s_rec = ParagraphStyle("R", parent=styles["Normal"], fontSize=10,
                textColor=rec_color, spaceBefore=2, spaceAfter=4, fontName="Helvetica-Bold")

            story.append(Paragraph(safe(r["name"]), s_h1))
            story.append(Paragraph(safe(r["recommendation"]), s_rec))
            story.append(Paragraph(f"Reason: {safe(r['brief_reason'])}", s_body))
            story.append(Paragraph(f"Main Risk: {safe(r['main_risk'])}", s_body))
            story.append(Spacer(1, 2*mm))
            story.append(Paragraph("Strengths", s_h2))
            for s in r["strengths"]:
                story.append(Paragraph(f"• {safe(s)}", s_body))
            story.append(Spacer(1, 2*mm))
            story.append(Paragraph("Risks", s_h2))
            for s in r["risks"]:
                story.append(Paragraph(f"• {safe(s)}", s_body))
            if r["validation_questions"]:
                story.append(Spacer(1, 2*mm))
                story.append(Paragraph("Validation Questions", s_h2))
                for i, q in enumerate(r["validation_questions"], 1):
                    story.append(Paragraph(f"{i}. {safe(q)}", s_body))
            if r["candidate_summary"]:
                story.append(Spacer(1, 2*mm))
                story.append(Paragraph("Candidate Summary", s_h2))
                story.append(Paragraph(safe(r["candidate_summary"]), s_body))
            story.append(Spacer(1, 4*mm))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#e0e0e0")))
            story.append(Spacer(1, 4*mm))

        doc.build(story)
        buffer.seek(0)
        return buffer

    pdf_buffer = generate_pdf(results)

    def generate_docx(results):
        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin = Mm(20)
            section.bottom_margin = Mm(20)
            section.left_margin = Mm(20)
            section.right_margin = Mm(20)

        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("Candidate Matching & Recommendation Report")
        run.bold = True
        run.font.size = Pt(16)

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub.add_run(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')} | Candidates: {len(results)}")
        sub_run.font.size = Pt(9)
        sub_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        doc.add_paragraph()

        # Recommendation summary heading
        h1 = doc.add_paragraph()
        h1_run = h1.add_run("Recommendation Summary")
        h1_run.bold = True
        h1_run.font.size = Pt(13)

        # Summary table
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, text in enumerate(["Candidate", "Recommendation", "Brief Reason", "Main Risk"]):
            p = hdr[i].paragraphs[0]
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            hdr[i]._tc.get_or_add_tcPr()

        REC_COLORS_DOCX = {
            "Move Forward": RGBColor(0x27, 0x67, 0x49),
            "Validate First": RGBColor(0x74, 0x42, 0x10),
            "Pass": RGBColor(0x74, 0x2A, 0x2A)
        }

        for r in results:
            row = table.add_row().cells
            for i, text in enumerate([r["name"], r["recommendation"], r["brief_reason"], r["main_risk"]]):
                p = row[i].paragraphs[0]
                run = p.add_run(text)
                run.font.size = Pt(8)
                if i == 1:
                    run.font.color.rgb = REC_COLORS_DOCX.get(r["recommendation"], RGBColor(0, 0, 0))
                    run.bold = True

        doc.add_paragraph()

        # Detailed reports
        h1 = doc.add_paragraph()
        h1_run = h1.add_run("Detailed Candidate Reports")
        h1_run.bold = True
        h1_run.font.size = Pt(13)

        for r in results:
            doc.add_paragraph()
            name_p = doc.add_paragraph()
            name_run = name_p.add_run(r["name"])
            name_run.bold = True
            name_run.font.size = Pt(11)

            rec_p = doc.add_paragraph()
            rec_run = rec_p.add_run(r["recommendation"])
            rec_run.bold = True
            rec_run.font.size = Pt(10)
            rec_run.font.color.rgb = REC_COLORS_DOCX.get(r["recommendation"], RGBColor(0, 0, 0))

            doc.add_paragraph(f"Reason: {r['brief_reason']}").runs[0].font.size = Pt(9)
            doc.add_paragraph(f"Main Risk: {r['main_risk']}").runs[0].font.size = Pt(9)

            s_h = doc.add_paragraph()
            s_h.add_run("Strengths").bold = True
            for s in r["strengths"]:
                p = doc.add_paragraph(f"• {s}", style="List Bullet")
                p.runs[0].font.size = Pt(9)

            r_h = doc.add_paragraph()
            r_h.add_run("Risks").bold = True
            for s in r["risks"]:
                p = doc.add_paragraph(f"• {s}", style="List Bullet")
                p.runs[0].font.size = Pt(9)

            if r["validation_questions"]:
                vq_h = doc.add_paragraph()
                vq_h.add_run("Validation Questions").bold = True
                for i, q in enumerate(r["validation_questions"], 1):
                    p = doc.add_paragraph(f"{i}. {q}")
                    p.runs[0].font.size = Pt(9)

            if r["candidate_summary"]:
                cs_h = doc.add_paragraph()
                cs_h.add_run("Candidate Summary").bold = True
                p = doc.add_paragraph(r["candidate_summary"])
                p.runs[0].font.size = Pt(9)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    docx_buffer = generate_docx(results)

    col1, col2 = st.columns(2)
    col1.download_button(
        label="Download PDF Report",
        data=pdf_buffer,
        file_name="candidate_report.pdf",
        mime="application/pdf"
    )
    col2.download_button(
        label="Download Word Report",
        data=docx_buffer,
        file_name="candidate_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
